import sqlite3
import pandas as pd
from docx import Document
from openpyxl import Workbook
from tkinter import messagebox

def save_tasks_to_file(conn, filename, file_format="txt"):
    try:
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM tasks ORDER BY priority ASC')
        tasks = cursor.fetchall()

        if file_format == "txt":
            with open(filename, 'w') as file:
                for task in tasks:
                    execution_date = task[4] if task[4] else ""
                    task_text = f"[{'x' if task[2] else ' '}] {task[1]} (Priorytet: {task[3]})"
                    if execution_date:
                        task_text += f", Data wykonania: {execution_date}"
                    file.write(task_text + '\n')
            messagebox.showinfo("Sukces", f"Zadania zapisane do pliku {filename} pomyślnie.")

        elif file_format == "xlsx":
            df = pd.DataFrame(tasks, columns=["id", "Zadanie", "Done", "Priorytet", "Data wykonania", "Notatka"])
            df["Data wykonania"] = df["Data wykonania"].apply(lambda x: "" if pd.isnull(x) else x)
            df.to_excel(filename, index=False)
            messagebox.showinfo("Sukces", f"Zadania zapisane do pliku {filename} pomyślnie.")

        elif file_format == "docx":
            doc = Document()
            doc.add_heading('Lista zadań', level=1)

            for task in tasks:
                execution_date = task[4] if task[4] else ""
                status = "x" if task[2] else " "
                task_text = f"[{status}] {task[1]} (Priorytet: {task[3]})"
                if execution_date:
                    task_text += f", Data wykonania: {execution_date}"
                doc.add_paragraph(task_text)

            doc.save(filename)
            messagebox.showinfo("Sukces", f"Zadania zapisane do pliku {filename} pomyślnie.")

        else:
            messagebox.showwarning("Uwaga", "Nieobsługiwany format pliku.")

    except sqlite3.Error as e:
        messagebox.showerror("Błąd SQLite", f"Błąd SQLite przy zapisywaniu zadań do pliku: {e}")
    except Exception as e:
        messagebox.showerror("Błąd", f"Błąd przy zapisie do pliku: {e}")